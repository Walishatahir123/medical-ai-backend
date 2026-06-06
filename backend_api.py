import os
import sys

os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["HF_DATASETS_OFFLINE"] = "1"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

from dotenv import load_dotenv
load_dotenv()

from docx import Document
from fastapi.responses import StreamingResponse
import io
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import sqlite3, hashlib

from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# =========================================================
# APP SETUP
# =========================================================

app = FastAPI(title="MedicalAI Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================================================
# DATABASE SETUP (AUTH)
# =========================================================

DB = "users.db"

def init_db():
    conn = sqlite3.connect(DB)
    conn.execute("""CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT, email TEXT UNIQUE, password TEXT)""")
    conn.commit()
    conn.close()

init_db()

def hash_pw(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

# =========================================================
# EMERGENCY DETECTION
# =========================================================

EMERGENCY_KEYWORDS = {
    "suicide": [
        "i want to die", "kill myself", "killing myself",
        "end my life", "suicide", "suicidal", "self harm",
        "self-harm", "hurt myself", "no reason to live",
        "don't want to live", "want to end it", "take my life",
        "better off dead", "can't go on", "give up on life",
    ],
    "heart_attack": [
        "chest pain", "chest pressure", "chest tightness",
        "left arm pain", "jaw pain", "heart attack",
        "shortness of breath", "heart is racing",
        "pain in chest", "crushing chest",
    ],
    "stroke": [
        "face drooping", "face numb", "slurred speech",
        "can't speak", "weakness on one side", "arm weakness",
        "sudden headache", "blurred vision", "stroke",
    ],
    "severe_bleeding": [
        "bleeding badly", "won't stop bleeding", "blood everywhere",
        "severe bleeding", "cut myself badly", "bleeding out",
    ],
    "unconscious": [
        "not breathing", "unconscious", "passed out",
        "not waking up", "fainted", "collapsed",
        "no pulse", "not responding",
    ],
}

EMERGENCY_MESSAGES = {
    "suicide": "🆘 YOU ARE NOT ALONE. Please call Umang helpline: 0311-7786264 or Rozan: 051-2890505. You matter. Please reach out to someone you trust immediately.",
    "heart_attack": "🚨 POSSIBLE HEART ATTACK. Call 1122 (Rescue) or 115 (Edhi Ambulance) immediately. Chew aspirin if available and sit down. Do NOT drive yourself.",
    "stroke": "🚨 POSSIBLE STROKE. Call 1122 NOW. Remember FAST: Face drooping, Arm weakness, Speech difficulty, Time to call emergency.",
    "severe_bleeding": "🚨 SEVERE BLEEDING. Call 1122 immediately. Apply firm pressure to the wound with a clean cloth and do not remove it.",
    "unconscious": "🚨 EMERGENCY. Call 1122 immediately. If not breathing, begin CPR if trained. Do not leave the person alone.",
}

def detect_emergency(text: str):
    text = text.lower()
    for category, keywords in EMERGENCY_KEYWORDS.items():
        for kw in keywords:
            if kw in text:
                return category
    return None

# =========================================================
# RAG SETUP
# =========================================================

DB_FAISS_PATH = "vectorstore/db_faiss"

print("Loading embedding model...")
embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True}
)

print("Loading FAISS database...")
try:
    faiss_db = FAISS.load_local(
        DB_FAISS_PATH,
        embedding_model,
        allow_dangerous_deserialization=True
    )
    retriever = faiss_db.as_retriever(search_type="similarity", search_kwargs={"k": 5})
    print("FAISS loaded successfully")
except Exception as e:
    print(f"WARNING: Could not load FAISS: {e}")
    faiss_db = None
    retriever = None

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    print("WARNING: GROQ_API_KEY not found in .env")

llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    model_name="llama-3.1-8b-instant",
    temperature=0.5,
    max_tokens=512
)

prompt = ChatPromptTemplate.from_template("""
You are a helpful medical assistant with access to two knowledge sources:

1. MEDICAL ENCYCLOPEDIA - Gale Encyclopedia of Medicine with detailed
   information about diseases, symptoms, treatments, and medications.

2. LAHORE DOCTOR DIRECTORY - A list of specialist doctors in Lahore,
   Pakistan, including their names, specializations, clinics, and
   contact information.

Use the context below to answer the user's question accurately.

Rules:
- If the user asks about a disease, symptom, treatment, or medication answer from the medical encyclopedia.
- If the user asks about a doctor, specialist, clinic, or hospital in Lahore answer from the doctor directory.
- If the user asks about both answer both parts.
- If the answer is not in the context, say: "I don't have enough information in my documents to answer that."
- Never make up doctor names, phone numbers, or addresses.

Context:
{context}

Question:
{input}

Answer:
""")

if retriever:
    combine_docs_chain = create_stuff_documents_chain(llm, prompt)
    rag_chain = create_retrieval_chain(retriever, combine_docs_chain)
    print("RAG chain ready")

# =========================================================
# AUTH MODELS & ROUTES
# =========================================================

class Register(BaseModel):
    name: str
    email: str
    password: str

class Login(BaseModel):
    email: str
    password: str

@app.post("/auth/register")
def register(r: Register):
    try:
        conn = sqlite3.connect(DB)
        conn.execute("INSERT INTO users (name,email,password) VALUES (?,?,?)",
                     (r.name, r.email, hash_pw(r.password)))
        conn.commit()
        conn.close()
        return {"message": "Registered successfully"}
    except Exception:
        raise HTTPException(status_code=400, detail="Email already exists")

@app.post("/auth/login")
def login(r: Login):
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    user = conn.execute("SELECT * FROM users WHERE email=? AND password=?",
                        (r.email, hash_pw(r.password))).fetchone()
    conn.close()
    if not user:
        raise HTTPException(status_code=400, detail="Invalid email or password")
    return {
        "token": "tok_" + str(user["id"]),
        "name": user["name"],
        "email": user["email"],
        "uid": user["id"]
    }


import smtplib
import secrets
from email.mime.text import MIMEText

class ForgotPassword(BaseModel):
    email: str

@app.post("/auth/forgot-password")
def forgot_password(r: ForgotPassword):
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    user = conn.execute("SELECT * FROM users WHERE email=?", (r.email,)).fetchone()
    conn.close()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")

    EMAIL_USER = os.getenv("EMAIL_USER")
    EMAIL_PASS = os.getenv("EMAIL_PASS")

    reset_link = f"http://localhost:5173/reset?email={r.email}"
    msg = MIMEText(f"Hello {user['name']},\n\nClick this link to reset your password:\n\n{reset_link}\n\nIf you did not request this, ignore this email.")
    msg["Subject"] = "MedicalAI - Password Reset Request"
    msg["From"] = EMAIL_USER
    msg["To"] = r.email

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(EMAIL_USER, EMAIL_PASS)
            server.send_message(msg)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Email failed: {str(e)}")

    return {"message": "Reset email sent"}
@app.post("/auth/logout")
def logout():
    return {"message": "Logged out"}
class GoogleAuth(BaseModel):
    email: str
    name: str
    google_id: str

@app.post("/auth/google")
def google_auth(r: GoogleAuth):
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    user = conn.execute("SELECT * FROM users WHERE email=?", (r.email,)).fetchone()
    if not user:
        conn.execute("INSERT INTO users (name,email,password) VALUES (?,?,?)",
                     (r.name, r.email, hash_pw(r.google_id)))
        conn.commit()
        user = conn.execute("SELECT * FROM users WHERE email=?", (r.email,)).fetchone()
    conn.close()
    return {
        "token": "tok_" + str(user["id"]),
        "name": user["name"],
        "email": user["email"],
        "uid": user["id"]
    }

# =========================================================
# CHAT ROUTES
# =========================================================

class ChatMessage(BaseModel):
    message: str
    session_id: str = "default"


# =========================================================
# ADMIN ROUTES
# =========================================================

ADMIN_EMAIL = "admin@medical.com"
ADMIN_PASSWORD = "admin123"

@app.post("/admin/login")
def admin_login(r: Login):
    if r.email != ADMIN_EMAIL or r.password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin credentials")
    return {"token": "admin_tok", "role": "admin", "name": "Admin"}

@app.get("/admin/users")
def get_users():
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    users = conn.execute("SELECT id, name, email FROM users").fetchall()
    conn.close()
    return {"users": [dict(u) for u in users], "total": len(users)}

@app.delete("/admin/users/{user_id}")
def delete_user(user_id: int):
    conn = sqlite3.connect(DB)
    conn.execute("DELETE FROM users WHERE id=?", (user_id,))
    conn.commit()
    conn.close()
    return {"message": "User deleted"}
@app.post("/chat/send")
def chat_send(body: ChatMessage):
    user_message = body.message

    # Emergency check first
    emergency = detect_emergency(user_message)
    if emergency:
        return {
            "answer": EMERGENCY_MESSAGES[emergency],
            "sources": []
        }

    if not retriever:
        raise HTTPException(status_code=503, detail="Vector store not loaded. Please rebuild FAISS index.")

    try:
        response = rag_chain.invoke({"input": user_message})
        answer = response["answer"]
        answer += "\n\n⚠️ *This information is for educational purposes only and does not replace professional medical advice.*"

        sources = []
        seen = set()
        for doc in response.get("context", []):
            source = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            key = f"{source}-{page}"
            if key not in seen:
                seen.add(key)
                sources.append({"source": source, "page": page})

        return {"answer": answer, "sources": sources}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat/save")
def chat_save(body: dict):
    return {"message": "Session saved"}

@app.get("/chat/history")
def chat_history():
    return []
# @app.get("/chat/download/{session_id}")
# def download_chat(session_id: str):
#     doc = Document()
#     doc.add_heading("MedicalAI Chat Session", 0)
#     doc.add_paragraph(f"Session ID: {session_id}")
#     doc.add_paragraph("This is your exported chat session.")
#     buf = io.BytesIO()
#     doc.save(buf)
#     buf.seek(0)
#     return StreamingResponse(
#         buf,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         headers={"Content-Disposition": f"attachment; filename=MedicalAI_{session_id}.docx"}
#     )
class DownloadRequest(BaseModel):
    messages: list = []

@app.post("/chat/download/{session_id}")
def download_chat(session_id: str, body: DownloadRequest):
    doc = Document()
    doc.add_heading("MedicalAI Chat Session", 0)
    doc.add_paragraph(f"Session ID: {session_id}")
    doc.add_paragraph("─" * 40)
    for msg in body.messages:
        role = "You" if msg.get("role") == "user" else "MedicalAI"
        time = msg.get("time", "")
        doc.add_paragraph(f"{role} [{time}]:", style="Heading 3" if msg.get("role") == "ai" else None)
        doc.add_paragraph(msg.get("text", ""))
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=MedicalAI_{session_id}.docx"}
    )

@app.get("/")
def root():
    return {"status": "MedicalAI backend running"}


